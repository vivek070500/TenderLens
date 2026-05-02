import os
import hashlib
import fitz  # PyMuPDF
import pdfplumber
from docx import Document as DocxDocument
from PIL import Image
import io

from modules.ocr import extract_text, extract_text_from_bytes


SUPPORTED_EXTENSIONS = {
    ".pdf", ".txt", ".docx", ".doc",
    ".jpg", ".jpeg", ".png", ".bmp", ".tiff",
}


def file_hash(filepath: str) -> str:
    """SHA-256 hash of a file for audit provenance."""
    h = hashlib.sha256()
    with open(filepath, "rb") as f:
        for chunk in iter(lambda: f.read(8192), b""):
            h.update(chunk)
    return h.hexdigest()


def detect_file_type(filepath: str) -> str:
    """Determine the file type from extension."""
    ext = os.path.splitext(filepath)[1].lower()
    if ext == ".pdf":
        return "pdf"
    elif ext == ".txt":
        return "text"
    elif ext in (".docx", ".doc"):
        return "docx"
    elif ext in (".jpg", ".jpeg", ".png", ".bmp", ".tiff"):
        return "image"
    return "unknown"


def _extract_pdf(filepath: str) -> list:
    """Extract text and tables from a PDF, falling back to OCR for scanned pages."""
    pages = []
    doc = fitz.open(filepath)

    for page_num in range(len(doc)):
        page = doc[page_num]
        text = page.get_text("text").strip()

        tables = []
        try:
            with pdfplumber.open(filepath) as pdf:
                if page_num < len(pdf.pages):
                    pp = pdf.pages[page_num]
                    for table in pp.extract_tables():
                        tables.append(table)
        except Exception:
            pass

        ocr_confidence = None
        is_scanned = len(text) < 50

        if is_scanned:
            pix = page.get_pixmap(dpi=300)
            img_bytes = pix.tobytes("png")
            ocr_result = extract_text_from_bytes(img_bytes)
            text = ocr_result["text"]
            ocr_confidence = ocr_result["confidence"]

        pages.append({
            "page_num": page_num + 1,
            "text": text,
            "tables": tables,
            "ocr_confidence": ocr_confidence,
            "is_scanned": is_scanned,
        })

    doc.close()
    return pages


def _extract_text_file(filepath: str) -> list:
    """Extract text from a plain text file."""
    with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
        content = f.read()
    return [{
        "page_num": 1,
        "text": content,
        "tables": [],
        "ocr_confidence": None,
        "is_scanned": False,
    }]


def _extract_docx(filepath: str) -> list:
    """Extract text from a Word document."""
    doc = DocxDocument(filepath)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    text = "\n".join(paragraphs)

    tables = []
    for table in doc.tables:
        rows = []
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        tables.append(rows)

    return [{
        "page_num": 1,
        "text": text,
        "tables": tables,
        "ocr_confidence": None,
        "is_scanned": False,
    }]


def _extract_image(filepath: str) -> list:
    """Extract text from an image via OCR."""
    image = Image.open(filepath)
    result = extract_text(image)
    return [{
        "page_num": 1,
        "text": result["text"],
        "tables": [],
        "ocr_confidence": result["confidence"],
        "is_scanned": True,
    }]


def ingest_document(filepath: str) -> dict:
    """
    Main entry point: parse any supported document and return structured data.
    """
    filename = os.path.basename(filepath)
    ftype = detect_file_type(filepath)
    fhash = file_hash(filepath)

    if ftype == "pdf":
        pages = _extract_pdf(filepath)
    elif ftype == "text":
        pages = _extract_text_file(filepath)
    elif ftype == "docx":
        pages = _extract_docx(filepath)
    elif ftype == "image":
        pages = _extract_image(filepath)
    else:
        pages = [{
            "page_num": 1,
            "text": "",
            "tables": [],
            "ocr_confidence": None,
            "is_scanned": False,
        }]

    full_text = "\n\n".join(p["text"] for p in pages if p["text"])

    min_ocr = None
    ocr_pages = [p["ocr_confidence"] for p in pages if p["ocr_confidence"] is not None]
    if ocr_pages:
        min_ocr = min(ocr_pages)

    return {
        "filename": filename,
        "filepath": filepath,
        "file_type": ftype,
        "file_hash": fhash,
        "pages": pages,
        "full_text": full_text,
        "page_count": len(pages),
        "min_ocr_confidence": min_ocr,
    }
